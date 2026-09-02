"""
Génère le modèle Word du papier en-tête CAP CONSEILS.

Le logo et la bande de couleurs vont dans l'en-tête de section, les mentions
légales dans le pied de page : Word les répète alors sur toutes les pages et
l'utilisateur ne peut pas les écraser en tapant son courrier.
"""
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x4E, 0x9C)
RED = RGBColor(0xC8, 0x1F, 0x3C)
INK = RGBColor(0x0A, 0x0F, 0x1C)
SLATE = RGBColor(0x4A, 0x55, 0x60)
MUTE = RGBColor(0x8B, 0x8F, 0x98)

BASE = "livrables/assets"


def shade(cell, hex_color):
    """Word n'a pas de `background-color` : la couleur d'un filet passe par
    l'ombrage de cellule d'un tableau sans bordure."""
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(el)


def fixed_layout(table, widths_mm):
    """Word répartit les colonnes lui-même tant que la table est en layout
    automatique. On force `fixed` et on écrit la grille : sans ça les filets
    de couleur ressortent tous à la même largeur."""
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    grid = table._tbl.find(qn("w:tblGrid"))
    for col, w in zip(grid.findall(qn("w:gridCol")), widths_mm):
        col.set(qn("w:w"), str(int(w * 56.7)))  # mm -> twips
    for row in table.rows:
        for cell, w in zip(row.cells, widths_mm):
            cell.width = Mm(w)


def no_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def color_bar(container, widths, colors, height_mm=3.2):
    """Bande de couleurs : un tableau d'une ligne, une cellule par bloc."""
    t = container.add_table(rows=1, cols=len(widths), width=Mm(sum(widths)))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    no_borders(t)
    fixed_layout(t, widths)
    for cell, c in zip(t.rows[0].cells, colors):
        shade(cell, c)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run()
        r.font.size = Pt(height_mm * 2.83)  # hauteur de ligne = hauteur du filet
    t.rows[0].height = Mm(height_mm)
    return t


def run(p, text, size=7.5, color=SLATE, bold=False, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    return r


doc = Document()

sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.left_margin = sec.right_margin = Mm(20)
sec.top_margin = Mm(16)
sec.bottom_margin = Mm(28)
sec.header_distance = Mm(10)
sec.footer_distance = Mm(8)

# ─────────────────────────────────────────────────────────── En-tête
hdr = sec.header
hdr.is_linked_to_previous = False
htab = hdr.add_table(rows=1, cols=2, width=Mm(170))
htab.autofit = False
no_borders(htab)
fixed_layout(htab, [60, 110])
left, right = htab.rows[0].cells

p = left.paragraphs[0]
p.paragraph_format.space_after = Pt(0)
p.add_run().add_picture(f"{BASE}/logo-full.png", width=Mm(34))

p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(4)
for txt, col in [("C", RED), ("omptabilité ", NAVY), ("A", RED), ("udit ", NAVY),
                 ("P", RED), ("atrimoine — ", NAVY), ("C", RED), ("onseils", NAVY)]:
    run(p, txt, size=9, color=col, bold=True)
p2 = right.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.paragraph_format.space_before = Pt(6)
p2.add_run().add_picture(f"{BASE}/label-co-pilotes.png", width=Mm(26))

bar = hdr.add_paragraph()
bar.paragraph_format.space_before = Pt(8)
bar.paragraph_format.space_after = Pt(0)
color_bar(hdr, [30, 15, 125], ["C81F3C", "2EC4F0", "1F4E9C"])

# ─────────────────────────────────────────────────────────── Pied de page
ftr = sec.footer
ftr.is_linked_to_previous = False
color_bar(ftr, [125, 15, 30], ["1F4E9C", "2EC4F0", "C81F3C"])

ftab = ftr.add_table(rows=1, cols=2, width=Mm(170))
ftab.autofit = False
no_borders(ftab)
fixed_layout(ftab, [152, 18])
fleft, fright = ftab.rows[0].cells

lines = [
    [("Cap Conseils Océan Indien", True), (" — SARL, société d'expertise comptable", False),
     ("  |  Capital 50 000 €  |  SIREN 489 543 660", False)],
    [("55, rue Estelle Darsanesing — 97420 Le Port  |  Tél. 0262 55 33 12  |  www.capconseils.net", False)],
    [("Inscrite au tableau de l'Ordre des experts-comptables — Conseil régional de La Réunion", False)],
]
first = True
for segs in lines:
    p = fleft.paragraphs[0] if first else fleft.add_paragraph()
    first = False
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for txt, bold in segs:
        run(p, txt, bold=bold, color=INK if bold else SLATE)

p = fleft.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(1)
run(p, "IBAN : FR76-1990-6009-7490-0137-7565-043", size=7, color=MUTE, italic=True)

p = fright.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run().add_picture(f"{BASE}/oec-mark.png", width=Mm(15))

# ─────────────────────────────────────────────────────────── Corps
def body(text, size=10.5, color=SLATE, align=None, before=0, after=6, indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent is not None:
        p.paragraph_format.left_indent = Mm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p

body("Destinataire\nRaison sociale\nAdresse\nCode postal — Ville",
     color=MUTE, indent=95, before=18, after=0)
body("Le Port, le [date]", color=MUTE, align=WD_ALIGN_PARAGRAPH.RIGHT, before=24, after=18)
body("Objet : […]", color=MUTE, after=12)
body("Madame, Monsieur,", color=MUTE, after=12)
body("[Corps du courrier]", color=MUTE)

doc.save("livrables/papier-entete.docx")
print("écrit : livrables/papier-entete.docx")
